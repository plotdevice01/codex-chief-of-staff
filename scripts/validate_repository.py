from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document

try:
    from .config_paths import ROOT
    from .live_acceptance_harness import validate_receipt
    from .test_persona import validate as validate_persona
    from .validate_icm import validate_icm
except ImportError:
    from config_paths import ROOT
    from live_acceptance_harness import validate_receipt
    from test_persona import validate as validate_persona
    from validate_icm import validate_icm


VERSION = (ROOT / "VERSION").read_text(encoding="utf-8").strip()
MANIFEST_VERSION = VERSION
REQUIRED = (
    ".gitattributes",
    ".agents/plugins/marketplace.json",
    ".codex-plugin/plugin.json",
    "plugin.json",
    "AGENTS.md",
    "CHANGELOG.md",
    "CONTEXT.md",
    "CONTRIBUTING.md",
    "LICENSE",
    "PRIVACY.md",
    "README.md",
    "SECURITY.md",
    "TERMS.md",
    "assets/icon.png",
    "assets/icon.svg",
    "assets/logo-dark.png",
    "assets/logo-dark.svg",
    "assets/logo.png",
    "assets/logo.svg",
    "assets/social-preview.png",
    "assets/social-preview.svg",
    "chief-of-staff.example.json",
    "docs/Codex Chief of Staff - Installation and SOP.docx",
    "docs/icm-conformance.md",
    "hooks/chief-of-staff-hook.js",
    "hooks/icm-enforcement-hook.js",
    "hooks/hooks.json",
    "install.ps1",
    "install.sh",
    "persona/Technical Assistant Persona - source.pdf",
    "persona/persona-contract.json",
    "persona/technical-assistant-persona.txt",
    "scripts/build_release.py",
    "scripts/build_sop.py",
    "scripts/configure.py",
    "scripts/live_acceptance_harness.py",
    "scripts/sync_project_agents.py",
    "scripts/test_persona.py",
    "scripts/validate_icm.py",
    "scripts/validate_install.py",
    "scripts/validate_local_parity.py",
    "skills/chief-of-staff/SKILL.md",
    "skills/chief-of-staff/agents/openai.yaml",
    "skills/chief-of-staff/internal/icm-architect/workflow.md",
    "skills/chief-of-staff/internal/icm-architect/LICENSE",
    "skills/chief-of-staff/internal/icm-architect/UPSTREAM.json",
    "skills/chief-of-staff/internal/icm-architect/agents/openai.yaml",
    "skills/chief-of-staff/references/universal-request-contract.md",
    "skills/chief-of-staff/references/capability-registry.json",
    "skills/chief-of-staff/references/content-production.md",
    "skills/chief-of-staff/references/live-acceptance.md",
    "skills/chief-of-staff/references/paid-video-creative.md",
    "skills/chief-of-staff/scripts/route_request.py",
    "skills/chief-of-staff/scripts/content_intelligence.py",
    "skills/chief-of-staff/scripts/validate_paid_video.py",
    "skills/chief-of-staff/vendor/manifest.json",
    "scripts/sync_content_runtime.py",
    "tests/test_hooks.js",
    "tests/test_content_runtime.py",
    "tests/test_live_acceptance_harness.py",
    "tests/fixtures/paid-video/valid-professional-service.json",
    "tests/fixtures/paid-video/sanitized-failed-campaign.json",
    "tests/test_icm.py",
    "tests/test_release.py",
    "tests/model-acceptance.json",
    "tests/receipts/chatgpt-work-v2.1.0.json",
    "tests/test_sync.py",
    "workflows/release/CONTEXT.md",
    "workflows/release/01_prepare/CONTEXT.md",
    "workflows/release/02_build/CONTEXT.md",
    "workflows/release/03_validate/CONTEXT.md",
    "workflows/release/04_publish/CONTEXT.md",
)
TEXT_SUFFIXES = {".json", ".md", ".ps1", ".py", ".sh", ".svg", ".txt", ".yaml", ".yml"}
PRIVATE_PATTERNS = {
    "Windows user path": re.compile(
        r"C:[\\/]+Users[\\/]+" + "Aa" + "ron", re.IGNORECASE
    ),
    "private Gmail identity": re.compile(
        "aaronz" + r"thomas@gmail\.com", re.IGNORECASE
    ),
    "private business identity": re.compile(
        "aaron@" + r"theingroup\.io", re.IGNORECASE
    ),
    "private client identity": re.compile(
        "aaron@" + r"drjonesdc\.com", re.IGNORECASE
    ),
    "private Slack user ID": re.compile(r"\bU08D9" + r"BUMQ95\b"),
    "private ClickUp workspace": re.compile(r"\b901414" + r"57186\b"),
    "private ClickUp user": re.compile(r"\b106" + r"489469\b"),
    "private project path": re.compile(
        "(" + "DEAN " + "DAYJOB|FAMILY " + "LAW|To Know " + "Oneself)",
        re.IGNORECASE,
    ),
}
PLACEHOLDER = re.compile(r"\b(TODO|TBD|REPLACE_WITH_[A-Z0-9_]+|YOUR_[A-Z0-9_]+)\b")
REPOSITORIES = (
    "https://github.com/plotdevice01/ai-sloppy-copy",
    "https://github.com/plotdevice01/brand-voice-factory",
    "https://github.com/plotdevice01/crafty-carousels-skill",
    "https://github.com/plotdevice01/codex-chief-of-staff",
)
INSTALL_VALUES = (
    "git clone https://github.com/plotdevice01/codex-chief-of-staff.git",
    r".\install.ps1",
    "./install.sh",
)
def read_json(relative: str) -> dict:
    return json.loads((ROOT / relative).read_text(encoding="utf-8-sig"))


def docx_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        ]
    )


def public_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return docx_text(path)
    return path.read_text(encoding="utf-8-sig")


def validate_install_guidance() -> list[str]:
    errors: list[str] = []
    for relative in ("README.md", "docs/installation.md", "docs/dependencies.md"):
        text = (ROOT / relative).read_text(encoding="utf-8-sig")
        for value in (*REPOSITORIES, *INSTALL_VALUES):
            if value not in text:
                errors.append(f"{relative} is missing install value: {value}")

    sop = ROOT / "docs/Codex Chief of Staff - Installation and SOP.docx"
    text = docx_text(sop)
    for command in INSTALL_VALUES:
        if command not in text:
            errors.append(f"SOP is missing install command: {command}")
    try:
        with zipfile.ZipFile(sop) as archive:
            relationships = archive.read("word/_rels/document.xml.rels").decode(
                "utf-8", errors="replace"
            )
    except (OSError, KeyError, zipfile.BadZipFile) as exc:
        errors.append(f"Could not read SOP hyperlinks: {exc}")
        relationships = ""
    for repository in REPOSITORIES:
        if repository not in relationships:
            errors.append(f"SOP is missing clickable repository link: {repository}")
    return errors


def validate_versions() -> list[str]:
    errors: list[str] = []
    release_values = {
        "chief-of-staff.example.json": read_json(
            "chief-of-staff.example.json"
        ).get("release_version"),
    }
    manifest_values = {
        "plugin.json": read_json("plugin.json").get("version"),
        ".codex-plugin/plugin.json": read_json(".codex-plugin/plugin.json").get(
            "version"
        ),
    }
    hook = (ROOT / "hooks/chief-of-staff-hook.js").read_text(encoding="utf-8")
    hook_match = re.search(r'const VERSION = "([^"]+)";', hook)
    release_values["hooks/chief-of-staff-hook.js"] = (
        hook_match.group(1) if hook_match else None
    )
    for name, value in release_values.items():
        if value != VERSION:
            errors.append(f"{name} version is {value!r}; expected {VERSION!r}.")
    for name, value in manifest_values.items():
        if value != MANIFEST_VERSION:
            errors.append(
                f"{name} host version is {value!r}; expected {MANIFEST_VERSION!r}."
            )
    return errors


def validate_model_acceptance(*, require_pass: bool = False) -> list[str]:
    errors: list[str] = []
    evidence = read_json("tests/model-acceptance.json")
    contract = read_json("persona/persona-contract.json")
    expected_tests = {
        item["id"] for item in contract.get("live_acceptance_tests", [])
    }
    if evidence.get("release_version") != VERSION:
        errors.append("Model acceptance evidence does not match the release version.")
    carried_from = evidence.get("carried_forward_from_release")
    receipt_version = carried_from if isinstance(carried_from, str) else VERSION
    work_receipt_path = (
        ROOT / "tests" / "receipts" / f"chatgpt-work-v{receipt_version}.json"
    )
    if work_receipt_path.exists():
        work_receipt = json.loads(work_receipt_path.read_text(encoding="utf-8"))
        errors.extend(
            f"ChatGPT Work receipt: {error}"
            for error in validate_receipt(
                work_receipt, "chatgpt-work", expected_version=receipt_version
            )
        )
    else:
        errors.append("ChatGPT Work release receipt is missing.")
    waived_models, waiver_errors = validate_release_waiver(evidence)
    errors.extend(waiver_errors)
    models = evidence.get("models", {})
    if set(models) != {"gpt-5.6-sol", "gpt-5.6-terra"}:
        errors.append("Model acceptance must cover GPT-5.6 Sol and GPT-5.6 Terra.")
        return errors
    for name, result in models.items():
        status = result.get("status")
        if status not in {"pass", "pending", "fail"}:
            errors.append(f"{name} model acceptance has invalid status {status!r}.")
        if status == "fail" and not result.get("evidence"):
            errors.append(f"{name} failed model acceptance needs evidence.")
        if require_pass and status != "pass" and name not in waived_models:
            errors.append(f"{name} model acceptance did not pass.")
        if set(result.get("tests", [])) != expected_tests:
            errors.append(f"{name} model acceptance does not cover every live test.")
    if models["gpt-5.6-sol"].get("reasoning_effort") != "medium":
        errors.append("GPT-5.6 Sol acceptance must use Medium reasoning.")
    for name, result in models.items():
        if str(result.get("evidence", "")).startswith("carried_forward_") and (
            result.get("model_facing_inputs_unchanged") is not True
            or not result.get("carried_forward_reason")
        ):
            errors.append(
                f"Carried-forward {name} evidence requires an unchanged-input "
                "declaration and reason."
            )
    if carried_from:
        if not re.fullmatch(r"\d+\.\d+\.\d+", str(carried_from)):
            errors.append("Carried-forward release must use a semantic version.")
        if evidence.get("patch_scope") != "documentation_and_distribution_only":
            errors.append("Carried-forward evidence requires a documentation-only patch scope.")
        for name, result in evidence.get("hosts", {}).items():
            if result.get("status") == "pass" and (
                not str(result.get("evidence", "")).startswith("carried_forward_")
                or not result.get("carried_forward_reason")
            ):
                errors.append(
                    f"Carried-forward {name} host evidence requires a reason."
                )
    hosts = evidence.get("hosts", {})
    if set(hosts) != {"codex", "chatgpt-work"}:
        errors.append("Model acceptance must cover Codex and ChatGPT Work hosts.")
    else:
        for name, result in hosts.items():
            status = result.get("status")
            if status not in {"pass", "pending", "fail"}:
                errors.append(f"{name} host acceptance has invalid status {status!r}.")
            if status == "fail" and not result.get("evidence"):
                errors.append(f"{name} failed host acceptance needs evidence.")
            if require_pass and status != "pass":
                errors.append(f"{name} host acceptance did not pass.")
        if hosts["chatgpt-work"].get("status") == "pass" and not work_receipt_path.exists():
            errors.append("ChatGPT Work cannot pass without its validated receipt.")
    smoke_status = evidence.get("installed_runtime_smoke", {}).get("status")
    if smoke_status not in {"pass", "pending", "fail"}:
        errors.append("Installed runtime smoke has an invalid status.")
    if require_pass and smoke_status != "pass":
        errors.append("Installed runtime smoke did not pass.")
    return errors


def validate_release_waiver(evidence: dict) -> tuple[set[str], list[str]]:
    waiver = evidence.get("release_waiver")
    if waiver is None:
        return set(), []

    errors: list[str] = []
    if waiver.get("status") != "approved":
        errors.append("Release waiver status must be approved.")
    if waiver.get("release_version") != VERSION:
        errors.append("Release waiver does not match the release version.")
    for field in ("approved_at", "approved_by", "reason"):
        if not isinstance(waiver.get(field), str) or not waiver[field].strip():
            errors.append(f"Release waiver requires {field}.")

    waived_checks = waiver.get("waived_checks")
    if not isinstance(waived_checks, list) or not waived_checks:
        errors.append("Release waiver requires at least one waived check.")
        return set(), errors
    if len(waived_checks) != len(set(waived_checks)):
        errors.append("Release waiver contains duplicate checks.")

    allowed = {"gpt-5.6-sol", "gpt-5.6-terra"}
    waived_models = {
        check.removeprefix("models.")
        for check in waived_checks
        if isinstance(check, str) and check.startswith("models.")
    }
    invalid = set(waived_checks) - {f"models.{name}" for name in allowed}
    if invalid:
        errors.append(
            "Release waiver may cover only pending Sol or Terra model checks."
        )
    models = evidence.get("models", {})
    for name in waived_models:
        if models.get(name, {}).get("status") != "pending":
            errors.append(f"Release waiver may cover only pending {name} evidence.")
    return waived_models & allowed, errors


def model_acceptance_release_status(evidence: dict) -> str:
    waived_models, waiver_errors = validate_release_waiver(evidence)
    if waiver_errors:
        return "candidate"
    if evidence.get("installed_runtime_smoke", {}).get("status") != "pass":
        return "candidate"
    if not all(
        result.get("status") == "pass"
        for result in evidence.get("hosts", {}).values()
    ):
        return "candidate"
    models = evidence.get("models", {})
    if not all(
        result.get("status") == "pass"
        or (result.get("status") == "pending" and name in waived_models)
        for name, result in models.items()
    ):
        return "candidate"
    return "pass_with_waiver" if waived_models else "pass"


def validate_manifest_paths() -> list[str]:
    errors: list[str] = []
    manifest = read_json(".codex-plugin/plugin.json")
    portable = read_json("plugin.json")
    portable_fields = {
        "$schema",
        "name",
        "version",
        "description",
        "author",
        "homepage",
        "repository",
        "license",
        "keywords",
        "extensions",
    }
    unknown_fields = sorted(set(portable) - portable_fields)
    if unknown_fields:
        errors.append(
            "Portable plugin manifest has nonconforming fields: "
            + ", ".join(unknown_fields)
        )
    if portable.get("$schema") != "https://agent-plugins.org/schemas/1.0.0/plugin.schema.json":
        errors.append("Portable plugin manifest has the wrong Agent Plugins schema.")
    portable_name = portable.get("name")
    if not isinstance(portable_name, str) or not re.fullmatch(
        r"(?!.*(?:--|\.\.))[a-z0-9](?:[a-z0-9.-]*[a-z0-9])?",
        portable_name,
    ) or len(portable_name) > 64:
        errors.append("Portable plugin manifest name violates Agent Plugins 1.0.0.")
    if portable.get("name") != manifest.get("name"):
        errors.append("Portable and Codex plugin names differ.")
    if portable.get("version") != manifest.get("version"):
        errors.append("Portable and Codex plugin versions differ.")
    discovered = sorted(ROOT.glob("skills/*/SKILL.md"))
    all_skill_entries = sorted(ROOT.glob("skills/**/SKILL.md"))
    expected_skill = ROOT / "skills" / "chief-of-staff" / "SKILL.md"
    if discovered != [expected_skill]:
        errors.append(
            "Agent Plugins discovery must expose exactly one skill: chief-of-staff."
        )
    if all_skill_entries != [expected_skill]:
        errors.append("Internal workflows must not contain nested SKILL.md entries.")
    for key, default in (("skills", None), ("hooks", "./hooks/hooks.json")):
        value = manifest.get(key, default)
        if not isinstance(value, str) or not (ROOT / value).exists():
            errors.append(f"Plugin manifest {key} path is missing: {value!r}.")
    interface = manifest.get("interface", {})
    for key in ("composerIcon", "logo", "logoDark"):
        value = interface.get(key)
        if not isinstance(value, str) or not (ROOT / value).is_file():
            errors.append(f"Plugin interface {key} path is missing: {value!r}.")
    marketplace = read_json(".agents/plugins/marketplace.json")
    plugins = marketplace.get("plugins", [])
    if marketplace.get("name") != "codex-chief-of-staff":
        errors.append("Agent Plugins marketplace name must be codex-chief-of-staff.")
    if len(plugins) != 1 or plugins[0].get("name") != "chief-of-staff":
        errors.append("Agent Plugins marketplace must expose exactly one Chief plugin.")
    elif plugins[0].get("source", {}).get("path") != "./":
        errors.append("Agent Plugins marketplace must install the repository root.")
    elif plugins[0].get("policy") != {
        "installation": "AVAILABLE",
        "authentication": "ON_INSTALL",
    }:
        errors.append("Agent Plugins marketplace must declare installation and authentication policy.")
    hooks = read_json("hooks/hooks.json")
    commands = [
        hook["command"]
        for event in hooks.get("hooks", {}).values()
        for group in event
        for hook in group.get("hooks", [])
    ]
    if not commands or any("PLUGIN_ROOT" not in command for command in commands):
        errors.append("Every Codex hook command must resolve PLUGIN_ROOT.")
    if any("CLAUDE_" in command for command in commands):
        errors.append("Codex hooks must not retain Claude-only environment variables.")
    event_names = set(hooks.get("hooks", {}))
    for event_name in ("UserPromptSubmit", "PreToolUse", "Stop"):
        if event_name not in event_names:
            errors.append(f"Shared hooks are missing the {event_name} ICM enforcement event.")
    if not any("icm-enforcement-hook.js" in command for command in commands):
        errors.append("Shared hooks do not load the ICM enforcement hook.")
    return errors


def validate_public_text() -> list[str]:
    errors: list[str] = []
    skip_placeholders = {
        "chief-of-staff.example.json",
        "scripts/validate_install.py",
        "scripts/validate_repository.py",
        "scripts/test_persona.py",
    }
    for path in sorted(ROOT.rglob("*")):
        if not path.is_file():
            continue
        relative = path.relative_to(ROOT).as_posix()
        if relative.startswith((".git/", "dist/", "qa/")):
            continue
        if path.suffix.lower() not in TEXT_SUFFIXES and path.suffix.lower() != ".docx":
            continue
        text = public_text(path)
        for label, pattern in PRIVATE_PATTERNS.items():
            if pattern.search(text):
                errors.append(f"{relative} contains {label}.")
        if relative not in skip_placeholders and PLACEHOLDER.search(text):
            errors.append(f"{relative} contains an unfinished placeholder.")
    return errors


def validate_archive(path: Path) -> list[str]:
    errors: list[str] = []
    try:
        with zipfile.ZipFile(path) as archive:
            names = [name.replace("\\", "/") for name in archive.namelist()]
    except (OSError, zipfile.BadZipFile) as exc:
        return [f"Invalid release archive: {exc}"]
    prefix = f"codex-chief-of-staff-v{VERSION}/"
    if not names or any(not name.startswith(prefix) for name in names):
        errors.append(f"Every archive entry must be under {prefix}.")
    if any(name.endswith("/chief-of-staff.json") for name in names):
        errors.append("Release archive contains private chief-of-staff.json.")
    for required in REQUIRED:
        if f"{prefix}{required}" not in names:
            errors.append(f"Release archive is missing {required}.")
    return errors


def validate_repository(
    archive: Path | None = None, *, require_model_acceptance: bool = False
) -> tuple[list[str], dict]:
    errors = [
        f"Missing required file: {name}" for name in REQUIRED if not (ROOT / name).is_file()
    ]
    errors.extend(validate_versions())
    errors.extend(validate_model_acceptance(require_pass=require_model_acceptance))
    errors.extend(validate_manifest_paths())
    errors.extend(validate_install_guidance())
    errors.extend(validate_public_text())
    icm_errors, _ = validate_icm()
    errors.extend(f"ICM validation: {error}" for error in icm_errors)
    persona_errors, metrics = validate_persona()
    errors.extend(f"Persona validation: {error}" for error in persona_errors)
    if archive:
        errors.extend(validate_archive(archive))
    return errors, metrics


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Validate the public Chief of Staff source and release archive."
    )
    parser.add_argument("--archive", type=Path)
    parser.add_argument("--require-model-acceptance", action="store_true")
    args = parser.parse_args()
    errors, metrics = validate_repository(
        args.archive, require_model_acceptance=args.require_model_acceptance
    )
    if errors:
        for error in errors:
            print(f"FAIL: {error}")
        return 1
    print(
        f"PASS: Chief of Staff v{VERSION} source is public-safe and version-consistent."
    )
    print(
        f"PASS: {metrics['persona_requirements']} persona requirements, "
        f"{metrics['integration_requirements']} integration rules, and "
        f"{metrics['live_acceptance_tests']} live scenarios preserved."
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
