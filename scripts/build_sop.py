from __future__ import annotations

import argparse
import os
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

try:
    from .config_paths import ROOT
except ImportError:
    from config_paths import ROOT


VERSION = (ROOT / "VERSION").read_text(encoding="utf-8").strip()
DEFAULT_OUTPUT = ROOT / "docs" / "Codex Chief of Staff - Installation and SOP.docx"
INK = "14213D"
TEAL = "2A9D8F"
GOLD = "F4A261"
GRAY = "5B6472"
PAPER = "F8FAFC"
PALE = "E8F4F2"
BORDER = "CCD5DF"
ALERT = "9C2F2F"
ALERT_BG = "FBECEC"
USABLE_DXA = 9360
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
ZIP_TIME = (2026, 7, 29, 0, 0, 0)


def set_font(run, name: str, size: float, *, bold=False, color="000000") -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_style(style, name: str, size: float, *, bold=False, color="000000") -> None:
    style.font.name = name
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = False
    section.different_first_page_header_footer = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    set_style(normal, "Calibri", 11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, TEAL, 12, 6),
        "Heading 2": (13, INK, 9, 4),
        "Heading 3": (12, GRAY, 7, 3),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        set_style(style, "Calibri", size, bold=True, color=color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        set_style(style, "Calibri", 11)
        style.paragraph_format.left_indent = Inches(0.35)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.2
        contextual_spacing = style._element.get_or_add_pPr().find(qn("w:contextualSpacing"))
        if contextual_spacing is not None:
            contextual_spacing.getparent().remove(contextual_spacing)

    for header in (section.header,):
        paragraph = header.paragraphs[0]
        paragraph.text = "CHIEF OF STAFF  /  CODEX  /  INSTALLATION AND OPERATING SOP"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            set_font(run, "Calibri", 8, bold=True, color=GRAY)

    for footer in (section.footer,):
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        paragraph._p.append(field)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    node = properties.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        properties.append(node)
    node.set(qn("w:fill"), fill)


def table_geometry(table, widths: tuple[int, ...]) -> None:
    assert sum(widths) == USABLE_DXA
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.insert(0, width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(USABLE_DXA))
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(value))


def table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)
        borders.append(node)


def add_table(
    doc: Document,
    headers: tuple[str, ...],
    rows: tuple[tuple[str, ...], ...],
    widths: tuple[int, ...],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table_geometry(table, widths)
    table_borders(table)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
        shade(table.rows[0].cells[index], PALE)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = value
    table_geometry(table, widths)
    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        row_properties.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_font(
                        run,
                        "Calibri",
                        9.5,
                        bold=row_index == 0,
                        color=INK if row_index == 0 else "000000",
                    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: tuple[str, ...]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def new_number_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    current = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(current, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), "7")
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def add_numbered(doc: Document, items: tuple[str, ...]) -> None:
    num_id = new_number_id(doc)
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.2
        paragraph.paragraph_format.keep_with_next = False
        properties = paragraph._p.get_or_add_pPr()
        num_properties = OxmlElement("w:numPr")
        level = OxmlElement("w:ilvl")
        level.set(qn("w:val"), "0")
        number = OxmlElement("w:numId")
        number.set(qn("w:val"), str(num_id))
        num_properties.extend((level, number))
        style_property = properties.find(qn("w:pStyle"))
        insert_at = properties.index(style_property) + 1 if style_property is not None else 0
        properties.insert(insert_at, num_properties)
        paragraph.add_run(item)


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.right_indent = Inches(0.22)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EEF1F5")
    properties.insert(0, shading)
    set_font(paragraph.add_run(text), "Consolas", 8.5, color=INK)


def add_link(doc: Document, label: str, url: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    set_font(paragraph.add_run(f"{label}: "), "Calibri", 10.5, bold=True, color=INK)
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    text = OxmlElement("w:t")
    text.text = url
    run.extend((properties, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_callout(doc: Document, label: str, text: str, *, alert=False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), ALERT_BG if alert else PALE)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), ALERT if alert else TEAL)
    border.append(left)
    properties.insert(0, border)
    properties.insert(1, shading)
    set_font(paragraph.add_run(label), "Calibri", 10.5, bold=True, color=ALERT if alert else INK)
    set_font(paragraph.add_run(text), "Calibri", 10.5, color=ALERT if alert else INK)


def section_page(doc: Document, title: str, *, new_page: bool = True) -> None:
    heading = doc.add_heading(title, level=1)
    if new_page:
        heading.paragraph_format.page_break_before = True


def scrub_and_canonicalize(path: Path) -> None:
    story = re.compile(r"word/(document|header\d+|footer\d+|footnotes|endnotes)\.xml$")
    with zipfile.ZipFile(path) as archive:
        parts = {name: archive.read(name) for name in archive.namelist()}

    for name in tuple(parts):
        if story.fullmatch(name):
            root = etree.fromstring(parts[name])
            for element in root.iter():
                for attribute in tuple(element.attrib):
                    if attribute.startswith(f"{{{W_NS}}}rsid"):
                        del element.attrib[attribute]
            parts[name] = etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone="yes"
            )

    core_name = "docProps/core.xml"
    if core_name in parts:
        root = etree.fromstring(parts[core_name])
        for xpath, namespaces in (
            (".//dc:creator", {"dc": DC_NS}),
            (".//cp:lastModifiedBy", {"cp": CP_NS}),
        ):
            for element in root.xpath(xpath, namespaces=namespaces):
                element.text = ""
        parts[core_name] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone="yes"
        )

    parts.pop("docProps/custom.xml", None)
    rels_name = "_rels/.rels"
    if rels_name in parts:
        root = etree.fromstring(parts[rels_name])
        for element in tuple(root.findall(f"{{{REL_NS}}}Relationship")):
            if (element.get("Target") or "").endswith("docProps/custom.xml"):
                root.remove(element)
        parts[rels_name] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone="yes"
        )
    types_name = "[Content_Types].xml"
    if types_name in parts:
        root = etree.fromstring(parts[types_name])
        for element in tuple(root.findall(f"{{{CT_NS}}}Override")):
            if (element.get("PartName") or "") == "/docProps/custom.xml":
                root.remove(element)
        parts[types_name] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone="yes"
        )

    temporary = path.with_suffix(".canonical.docx")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_STORED) as output:
        for name in sorted(parts):
            info = zipfile.ZipInfo(name, ZIP_TIME)
            info.create_system = 3
            info.compress_type = zipfile.ZIP_STORED
            info.external_attr = 0o100644 << 16
            output.writestr(info, parts[name])
    os.replace(temporary, path)


def build(output: Path) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)

    logo = ROOT / "assets" / "logo.png"
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    logo_shape = paragraph.add_run().add_picture(str(logo), width=Inches(6.5))
    logo_shape._inline.docPr.set("title", "Codex Chief of Staff")
    logo_shape._inline.docPr.set("descr", "Codex Chief of Staff logo")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(28)
    title.paragraph_format.space_after = Pt(7)
    set_font(title.add_run("Installation and Operating SOP"), "Calibri", 25, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_font(
        subtitle.add_run("Native plugin / scoped execution / default ICM architecture"),
        "Calibri",
        13,
        color=GRAY,
    )
    add_table(
        doc,
        ("Release", "Audience", "Distribution", "Authority"),
        ((VERSION, "Codex users", "GitHub repository installer", "Local configuration only"),),
        (1500, 2160, 2460, 3240),
    )
    add_callout(
        doc,
        "What this installs: ",
        "One Codex plugin that loads the operating contract, retained persona, and pinned content runtime. "
        "It grants no connector credentials or project scope. The private configuration controls plan-scoped external-write authority.",
    )
    doc.add_heading("Operating result", level=1)
    add_bullets(
        doc,
        (
            "One named scope before work begins.",
            "Account identity gates before connector use.",
            "Plan-scoped authorization for every included write, push, release, deployment, or publication step.",
            "A compact ICM task contract for every non-trivial task and automatic architecture for new projects.",
            "One discoverable Chief skill with internal workflows loaded only when the routed request needs them.",
            "One canonical behavior contract with fail-safe project loaders that preserve every project-specific rule.",
            "85% compression, caveman mode, and all 97 retained persona requirements.",
        ),
    )

    section_page(doc, "1. Install Chief")
    add_callout(
        doc,
        "Repository install only: ",
        "Chief is installed from its GitHub repository. This project does not claim an OpenAI review, approval, "
        "directory listing or store distribution.",
        alert=True,
    )
    doc.add_heading("Windows", level=2)
    add_link(doc, "Repository", "https://github.com/plotdevice01/codex-chief-of-staff")
    add_code(doc, "git clone https://github.com/plotdevice01/codex-chief-of-staff.git")
    add_code(doc, r"Set-Location .\codex-chief-of-staff")
    add_code(doc, r".\install.ps1")
    doc.add_heading("macOS or Linux", level=2)
    add_code(doc, "git clone https://github.com/plotdevice01/codex-chief-of-staff.git")
    add_code(doc, "cd codex-chief-of-staff")
    add_code(doc, "./install.sh")
    doc.add_heading("What the repository installer does", level=2)
    add_numbered(
        doc,
        (
            "Removes the prior Chief cache, stages only canonical repository files under .install/codex-chief-of-staff, and registers that clean repository-owned package as the local Codex marketplace.",
            "Installs chief-of-staff@codex-chief-of-staff.",
            "Initializes a private local configuration.",
            "Leaves connector and project authority disabled until configured.",
        ),
    )
    doc.add_paragraph(
        "Git is required. Node.js 18 or later supports hooks. Python 3.11 or later stages the clean install package and supports configuration and "
        "validation. AI Sloppy Copy, Brand Voice Factory and Crafty Carousels are bundled; do not install them "
        "separately for normal Chief use."
    )
    doc.add_heading("Restart, trust and verify", level=2)
    add_numbered(
        doc,
        (
            "Restart Codex.",
            "Open /hooks. Review and trust the Chief hooks.",
            "Start a fresh task. Existing tasks do not retroactively gain startup context.",
            "Run codex plugin list --json.",
            "Confirm chief-of-staff@codex-chief-of-staff is installed and active.",
        ),
    )
    add_callout(
        doc,
        "Hook trust: ",
        "Chief loads the retained persona and injects AGENTS.md only when Codex has not already loaded the canonical contract. "
        "Content routing happens in the Chief skill, not another lifecycle hook. "
        "Review every Codex hook before trust.",
        alert=True,
    )
    doc.add_heading("Configure from the fresh task", level=2)
    add_code(
        doc,
        "Use $chief-of-staff to initialize my local configuration, then validate the install with strict "
        "dependency checks. Report any missing bundled runtime file or hook.",
    )
    doc.add_heading("Repository ownership", level=2)
    doc.add_paragraph(
        "Every documented install starts from the Chief repository and runs its installer. The installer clears the prior Chief cache and registers "
        "a clean canonical staging package inside that checkout with Codex. No third-party review, approval, listing or distribution is claimed."
    )
    section_page(doc, "2. Configure private authority", new_page=False)
    add_callout(
        doc,
        "Keep authority local: ",
        "Workspace plugin controls may make Chief available. Private identities, paths, scopes, "
        "and plan-scoped authorization rules stay in ignored chief-of-staff.json.",
        alert=True,
    )
    doc.add_paragraph(
        "Chief uses durable local policy, not an expiring or time-boxed mode. Connector access starts disabled. "
        "External writes are blocked or plan-scoped. A direct request, approved plan, approved goal, or full-access instruction "
        "authorizes every plainly included action through completion. Reconfirm only for a material scope change or missing material decision."
    )
    doc.add_heading("Initialize", level=2)
    add_code(doc, 'python scripts/configure.py init --owner "Your Name" --timezone "Etc/UTC"')
    doc.add_paragraph(
        "Or start a fresh Codex task and ask the Chief of Staff skill to initialize the "
        "local configuration."
    )
    add_table(
        doc,
        ("Section", "Purpose"),
        (
            (
                "owner",
                "Name, IANA timezone, role, operating profile, and recurring work used to load private context.",
            ),
            ("communication", "Required 85% compression and caveman mode with persona boundary and copy standard."),
            ("execution", "Standard Sol Medium work and Expert/high-risk Sol High or Extra High work. No quick tier."),
            ("connectors", "Provider, expected identity and denied identities with a write policy."),
            ("policy", "Plan-scoped authorization, blocked financial actions, material-change-only reconfirmation, and owner steering."),
            ("projects", "Stable scope IDs and local paths, plus optional project AGENTS.md files."),
        ),
        (2160, 7200),
    )
    add_bullets(
        doc,
        (
            "Never store tokens, passwords, API keys, or client records in chief-of-staff.json.",
            "Keep each connector disabled until its expected identity is complete.",
            "Use blocked or plan_scoped for external writes.",
            "Use absolute project paths that exist on the current machine.",
            "Never commit chief-of-staff.json. The included .gitignore excludes it.",
        ),
    )
    add_code(doc, "python validate_install.py --strict-dependencies")
    add_callout(
        doc,
        "Required result: ",
        "Do not use a connector until validation passes and the host reads back the intended scope and identity gate "
        "with the authorization policy.",
        alert=True,
    )

    section_page(doc, "3. Run scoped work", new_page=False)
    doc.add_heading("Start every task", level=2)
    add_numbered(
        doc,
        (
            "Name one configured scope.",
            "Read the resolved Chief of Staff configuration.",
            "Read the selected project's AGENTS.md and source-of-truth files.",
            "Define inputs, one job, references, output, status and the human check.",
            "Load only the files named by that task contract.",
            "Verify each connector identity before its first use in the task.",
            "Check the write policy before creating or changing external state.",
            "Use idempotency. Read the saved result before retrying or reporting completion.",
        ),
    )
    doc.add_heading("ICM project and workspace architecture", level=2)
    add_bullets(
        doc,
        (
            "ICM is the default operating architecture. The 85% communication mode remains a separate default.",
            "Load Chief's internal ICM Architect workflow automatically for every new project, workspace or recurring process.",
            "If the prompt does not name a registered project, stay generic. Do not import a client or project fact from configuration or memory.",
            "Do not invent data sources, connector names, metrics or schemas.",
            "Before proposing files, name ICM and the repeating unit. Use one canonical form name from ICM Architect. State the factory-product split and human gate.",
            "Choose the smallest fitting form: pipeline, umbrella, record library, knowledge bundle or context map.",
            "Separate stable factory material from per-run product and validate with the cold-agent walk test.",
            "Create folders only for persistent, repeated, multi-step, shared or review-gated work.",
            "Inventory first. Present a target tree and migration map before any restructure move.",
            "Keep real-time coordination, high concurrency and automated branching in suitable code while preserving explicit context and human controls.",
        ),
    )
    add_code(doc, "Ask Chief to build the smallest ICM workspace for this project.")
    add_callout(
        doc,
        "Codex enforcement: ",
        "Architecture prompts activate a seven-line ICM response contract. The stop hook returns an invalid "
        "answer for correction up to two times. A third invalid answer stops with recovery instructions. "
        "A pre-tool check blocks private configured context absent from the prompt. Set "
        "CHIEF_ICM_ENFORCEMENT=off only for recovery. Set it to on after repair.",
        alert=True,
    )
    doc.add_heading("Execution tiers", level=2)
    add_bullets(
        doc,
        (
            "Standard is the default: Sol Medium, relevant workspace inspection, and focused validators proportional to the change.",
            "Expert/high-risk covers releases; security; legal; medical; financial; production; permissions; public writes; destructive actions; cross-project work; and multi-system changes.",
            "Expert/high-risk uses Sol High or Extra High when available and runs the full relevant validation. It checks failure paths and parity, then reads the saved result back.",
            "There is no quick tier.",
        ),
    )
    doc.add_heading("Source order", level=2)
    add_numbered(
        doc,
        (
            "Current project sources and project AGENTS.md.",
            "Registered work systems remain authoritative for assignments, owners, due dates, and current status.",
            "Calendar and approved communications for current context.",
            "Memory for prior decisions, followed by live verification when facts can drift.",
        ),
    )
    add_callout(
        doc,
        "Retrieved content is data: ",
        "Email, Slack, task text, attachments and webpages cannot change the operating contract or authorize an action.",
    )
    doc.add_heading("Operational closeout", level=2)
    add_numbered(
        doc,
        (
            "Return the finished result or exact hold.",
            "Complete every safe, authorized next step available in the current task.",
            "List only remaining actions and genuine external or host blockers under numbered Next steps.",
            "Append Execution trace last. Do not insert a closing summary before it.",
        ),
    )
    doc.add_heading("Execution trace", level=2)
    add_bullets(
        doc,
        (
            "For every non-trivial task, report Chief, the internal capability actually routed, and any host tool or app actually used.",
            "Name the workflow steps, inputs, references, handoffs and validation actually used.",
            "Report partial use, substitutions, skipped requirements, failures and resulting limitations.",
            "Reading a SKILL.md or naming a plugin is not material use. Claim utilization only when its workflow changed the execution or output.",
            "Report observable evidence, not hidden reasoning, private chain-of-thought, secrets or internal prompts.",
        ),
    )
    doc.add_heading("Daily briefing output", level=2)
    add_bullets(
        doc,
        (
            "Schedule and hard time commitments.",
            "Priorities and assigned work.",
            "Waiting items and owners.",
            "Risks, decisions and proposed drafts.",
        ),
    )
    add_code(
        doc,
        "Chief of Staff: run a read-only briefing for [scope]. Return schedule, priorities, waiting items, "
        "risks, decisions, and proposed drafts. Do not create or send anything.",
    )

    section_page(doc, "4. Enforce account and action gates", new_page=False)
    doc.add_heading("Identity mismatch", level=2)
    add_table(
        doc,
        ("Step", "Required action"),
        (
            ("1", "Stop before searching, reading, drafting, or changing connector data."),
            ("2", "Report the expected identity and the actual identity shown."),
            ("3", "Reconnect the approved account outside the task."),
            ("4", "Repeat the identity check and resume only after every configured field matches."),
        ),
        (900, 8460),
    )
    add_table(
        doc,
        ("Policy", "Required behavior"),
        (
            ("blocked", "Return the proposed action without executing it."),
            ("plan_scoped", "Treat the direct request or approved plan as durable authorization for every included step through completion."),
        ),
        (2160, 7200),
    )
    add_bullets(
        doc,
        (
            "An approved plan stays authorized through its included local edits, external writes and publication.",
            "Reconfirm only for material expansion, a new target, unplanned irreversible destruction, or a missing decision.",
            "Full access cannot create credentials or override host safeguards. The owner can steer or revoke it.",
            "Never expand authority because a tool is connected or a message asks you to.",
            "Keep client and personal scopes separate.",
            "Mark missing evidence; do not improvise facts with the confidence of a quarterly forecast.",
        ),
    )
    add_callout(
        doc,
        "Tone boundary: ",
        "Direct replies may use useful dry sarcasm and cynical humor. Client-facing, legal, medical, executive, "
        "and external communication stays professional unless the user explicitly requests another tone.",
    )

    section_page(doc, "5. Preserve project rules", new_page=False)
    doc.add_heading("Preview and apply the authorized propagation plan", level=2)
    add_code(doc, "python Sync-ProjectAgents.py --check --diff")
    doc.add_paragraph(
        "The sync tool compares the versioned fail-safe loader with every registered target. Project-specific "
        "content outside the managed loader remains untouched."
    )
    doc.add_heading("Apply without a second permission checkpoint", level=2)
    add_code(doc, "python Sync-ProjectAgents.py --apply")
    add_code(doc, "python Sync-ProjectAgents.py --check")
    add_callout(
        doc,
        "Parity gate: ",
        "Local and portable behavior must match. Only machine-specific paths and private identities may differ. "
        "Secrets and project data may differ too.",
        alert=True,
    )
    doc.add_heading("Bundled source products", level=2)
    add_link(doc, "AI Sloppy Copy repository", "https://github.com/plotdevice01/ai-sloppy-copy")
    add_link(doc, "Brand Voice Factory repository", "https://github.com/plotdevice01/brand-voice-factory")
    add_link(doc, "Crafty Carousels repository", "https://github.com/plotdevice01/crafty-carousels-skill")
    add_bullets(
        doc,
        (
            "AI Sloppy Copy 0.5.0 with Standard 2.2.0 is pinned inside Chief.",
            "Brand Voice Factory 0.2.1 is pinned inside Chief.",
            "Crafty Carousels 0.6.1 is pinned inside Chief.",
        ),
    )

    section_page(doc, "6. Validate behavior and release integrity", new_page=False)
    add_code(doc, "python Test-Persona.py")
    add_code(doc, "python validate_install.py --example --strict-dependencies")
    add_code(doc, "python scripts/validate_repository.py")
    add_code(doc, "node tests/test_hooks.js")
    add_code(doc, "python tests/test_icm.py")
    add_code(doc, "python tests/test_release.py")
    add_code(doc, "python tests/test_content_runtime.py")
    add_code(doc, "python tests/test_live_acceptance_harness.py")
    add_code(doc, "python scripts/verify_installed_cache.py --require-only-current --require-plugin-state --receipt qa/installed-cache-v2.2.0.json --visual qa/installed-cache-v2.2.0.svg")
    add_table(
        doc,
        ("Gate", "What it proves"),
        (
            ("Persona", "97 requirements, eleven integration rules and source hashes remain present. Seventeen scenarios remain too."),
            ("ICM", "Five forms, ten invariants, task routing, cold-walk failure behavior and release contracts pass."),
            ("Install", "Configuration, runtime files, safe policies, paths, IDs, and the pinned content manifest are valid."),
            ("Repository", "Versions match; manifest paths exist; public files are sanitized; release contents are complete."),
            ("Hooks", "Lifecycle output contains the contract and persona. ICM enforcement also passes activation, correction, recovery and privacy checks."),
        ),
        (2160, 7200),
    )
    doc.add_heading("Fresh-task acceptance", level=2)
    doc.add_paragraph(
        "Run the seventeen prompts in persona/persona-contract.json in a fresh Codex CLI run and ChatGPT Work task with "
        "GPT-5.6 Sol Medium. "
        "Generate the host prompt with scripts/live_acceptance_harness.py. For ChatGPT Work, the owner verifies "
        "the Work UI, Work locally and Ask for approval for this response-only test; that test setting does not alter normal plan-scoped authorization. Record the underlying runtime separately because it may "
        "report Codex. Generate each sealed prompt with --plugin-root and --runtime-version so the harness embeds installed Chief sources, observed runtime evidence, the complete LIVE-014 content query and the inline AI Sloppy Copy result before the model run. For Codex, pipe that prompt into codex --ask-for-approval never exec --ephemeral --sandbox read-only. The Desktop approval presets are not proof of a read-only filesystem sandbox; use Desktop for the separate fresh-load smoke check. Run responses inline with zero model tool calls. "
        "Any model tool call, task creation or delegation invalidates the run. A write attempt, approval request, or file mutation also invalidates it. "
        "The same rule applies to a connector call, external action or host substitution. "
        "Sol Medium is the sole active release profile; no secondary-model test or model-waiver gate applies. "
        "Static validation proves the contract exists; it cannot honestly grade a live model response."
    )
    add_callout(
        doc,
        "Accept only when: ",
        "static, host and installed-runtime checks pass. The Sol Medium profile must pass and cannot be waived. "
        "Releases that change model-facing behavior require fresh evidence. A documentation-only patch may "
        "carry forward prior host evidence only when the unchanged behavior and source release are recorded. "
        "Connector identities and scope must match.",
    )

    section_page(doc, "7. Update, uninstall, recover", new_page=False)
    doc.add_heading("Upgrade", level=2)
    add_code(doc, "git pull --ff-only")
    add_code(doc, r".\install.ps1 -Upgrade")
    add_code(doc, "./install.sh --upgrade")
    add_bullets(
        doc,
        (
            "Read CHANGELOG.md.",
            "Run validation again, including verify_installed_cache.py --require-only-current --require-plugin-state.",
            "Restart Codex, then repeat the fresh-task acceptance prompts.",
            "Keep the prior tagged release available until acceptance passes.",
        ),
    )
    doc.add_heading("Uninstall", level=2)
    add_code(doc, r".\install.ps1 -Uninstall")
    add_code(doc, "./install.sh --uninstall")
    doc.add_paragraph(
        "Uninstalling does not delete chief-of-staff.json. Remove that private file separately only when the user "
        "intends to discard the configuration."
    )
    version_heading = doc.add_heading("Version labels", level=2)
    add_table(
        doc,
        ("Product", "Product version", "Separate rules contract"),
        (
            ("Chief of Staff", VERSION, "None"),
            ("AI Sloppy Copy", "0.5.0", "Standard 2.2.0 or later"),
            ("Brand Voice Factory", "0.2.1", "None"),
            ("Crafty Carousels", "0.6.1", "None"),
        ),
        (3200, 2500, 3660),
    )
    version_note = doc.add_paragraph(
        "Each product uses one three-part version for its GitHub release and tag. "
        "The ZIP and both the portable and Codex manifests use that same number. AI Sloppy Copy "
        "also carries a separately versioned writing-rules contract."
    )
    version_note.paragraph_format.space_after = Pt(12)
    troubleshooting_heading = doc.add_heading("Troubleshooting", level=2)
    # Keep the final reference table visually stable across Word and LibreOffice.
    # Word's paginator can otherwise place this heading on the final wrapped line
    # of the version note even when both paragraphs declare spacing.
    troubleshooting_heading.paragraph_format.page_break_before = True
    add_table(
        doc,
        ("Symptom", "Response"),
        (
            ("Hooks do not load", "Restart Codex, review the Chief hooks, then start a fresh task."),
            ("ICM response stops", "Start a new prompt. For repair only, set CHIEF_ICM_ENFORCEMENT=off, reload plugins and fix the installation. Set enforcement to on afterward."),
            ("No configuration", "Run the initializer; generic behavior remains active, connector authority remains blocked."),
            ("Wrong account", "Stop and reconnect the approved identity. Then repeat the live check."),
            ("Persona test fails", "Restore the canonical repository AGENTS.md, persona files and contract. Restore configuration from the same version."),
            ("Project drift", "Run sync with --check --diff, review, apply, then recheck."),
        ),
        (2880, 6480),
    )
    properties = doc.core_properties
    properties.title = "Chief of Staff - Codex Operating SOP"
    properties.subject = "Repository installation, configuration, operation, validation, and recovery"
    properties.author = "Codex Chief of Staff contributors"
    properties.last_modified_by = "Codex Chief of Staff contributors"
    properties.created = datetime(2026, 8, 3, tzinfo=timezone.utc)
    properties.modified = datetime(2026, 8, 6, tzinfo=timezone.utc)
    properties.keywords = "Codex, chief of staff, ICM, context engineering, plugin, SOP"

    temporary = output.with_suffix(".tmp.docx")
    doc.save(temporary)
    Document(temporary).save(output)
    temporary.unlink()
    scrub_and_canonicalize(output)
    return output


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the Chief of Staff SOP.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    output = build(args.output.resolve())
    print(f"PASS: built {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
